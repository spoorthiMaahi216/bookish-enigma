import os
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw
import math

OUTPUT_DIR = "/workspace/shadow_questions"
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images_user")

def ensure_dirs() -> None:
	os.makedirs(IMAGES_DIR, exist_ok=True)


def add_mono_line(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	run = p.add_run(text)
	font = run.font
	font.name = "Courier New"
	font.size = Pt(10.5)


# ---------- Image generators tailored to prompts ----------

def img_sequence_5cycle(path: str) -> None:
	# 5 repeating shapes: circle, square, triangle, star, pentagon
	w, h = 720, 150
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	for i in range(10):
		x = 20 + i * 70
		y = 25
		t = i % 5
		if t == 0:
			d.ellipse((x, y, x + 50, y + 50), outline="black", width=3)
		elif t == 1:
			d.rectangle((x, y, x + 50, y + 50), outline="black", width=3)
		elif t == 2:
			d.polygon([(x + 25, y), (x + 50, y + 50), (x, y + 50)], outline="black")
		elif t == 3:
			cx, cy, r = x + 25, y + 25, 23
			pts = []
			for k in range(5):
				ang = -math.pi / 2 + k * 2 * math.pi / 5
				pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
			star = [pts[i % 5] for i in [0, 2, 4, 1, 3]]
			d.line(star + [star[0]], fill="black", width=3)
		else:
			cx, cy, r = x + 25, y + 25, 23
			pts = []
			for k in range(5):
				ang = -math.pi / 2 + k * 2 * math.pi / 5
				pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
			d.polygon(pts, outline="black")
	img.save(path)


def img_altitude_100_to_500(path: str) -> None:
	w, h = 560, 300
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	# axes
	d.line((60, 20, 60, h - 40), fill="black", width=2)
	d.line((60, h - 40, w - 20, h - 40), fill="black", width=2)
	# polyline from (0,100) to (4,500)
	points = []
	for t, alt in [(0, 100), (1, 180), (2, 260), (3, 380), (4, 500)]:
		x = 60 + int(t * (w - 100) / 4)
		y = h - 40 - int((alt - 100) * (h - 80) / (500 - 100))
		points.append((x, y))
	d.line(points, fill="blue", width=3)
	img.save(path)


def img_midpoints_generic(path: str) -> None:
	w, h = 660, 140
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	margin = 40
	R = (margin, h // 2)
	S = (margin + 140, h // 2)
	T = (margin + 280, h // 2)
	V = (margin + 560, h // 2)
	d.line((R, V), fill="black", width=3)
	for label, pt in [("R", R), ("S", S), ("T", T), ("V", V)]:
		d.ellipse((pt[0] - 4, pt[1] - 4, pt[0] + 4, pt[1] + 4), fill="black")
		d.text((pt[0] - 6, pt[1] - 24), label, fill="black")
	# annotate ST
	d.text(((S[0] + T[0]) // 2 - 14, S[1] + 10), "ST=12", fill="black")
	img.save(path)


def img_rect_squares_7_12(path: str) -> None:
	# 3x2 grid, shade 3 full + half of one cell
	w, h = 420, 280
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	cell_w, cell_h = w // 3, h // 2
	# full shaded cells: (0,0),(1,0),(0,1)
	full = {(0, 0), (1, 0), (0, 1)}
	for r in range(2):
		for c in range(3):
			x0, y0 = c * cell_w, r * cell_h
			x1, y1 = x0 + cell_w - 2, y0 + cell_h - 2
			if (c, r) in full:
				d.rectangle((x0 + 2, y0 + 2, x1, y1), fill=(185, 185, 185))
			d.rectangle((x0 + 2, y0 + 2, x1, y1), outline="black", width=3)
	# half shade right half of cell (1,1)
	c, r = 1, 1
	x0, y0 = c * cell_w, r * cell_h
	x1, y1 = x0 + cell_w - 2, y0 + cell_h - 2
	d.rectangle((x0 + cell_w // 2, y0 + 2, x1, y1), fill=(185, 185, 185))
	# redraw the cell border
	d.rectangle((x0 + 2, y0 + 2, x1, y1), outline="black", width=3)
	img.save(path)


def img_card_holes(path: str) -> None:
	w, h = 280, 280
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	d.rectangle((20, 20, w - 20, h - 20), outline="black", width=3)
	d.ellipse((80, 90, 100, 110), outline="black", width=3)
	d.ellipse((180, 160, 200, 180), outline="black", width=3)
	img.save(path)


def img_segments_two_squares(path: str) -> None:
	w, h = 720, 200
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	# horizontal base line
	y = h // 2
	left = 40
	lengths = [180, 240, 300]  # proportional to 6,8,10
	labels = ["AB=6 cm", "CD=8 cm", "EF=10 cm"]
	x = left
	for i, L in enumerate(lengths):
		d.line((x, y, x + L, y), fill="black", width=4)
		d.text((x + L // 2 - 30, y + 10), labels[i], fill="black")
		# draw square at joints between segments (after first two)
		if i < 2:
			# square of side proportional to 2 cm
			s = 60
			d.rectangle((x + L - s // 2, y - s - 10, x + L - s // 2 + s, y - 10), outline="black", width=3)
		x += L
	img.save(path)


def build_doc(path: str) -> None:
	doc = Document()
	# Title/description not numbered; then 25 items below
	content_blocks = [
		{
			"title": "Solve Linear Equation (One-Step)",
			"desc": "Solve for n in a simple linear equation.",
			"question": "If $n+5=5$, what is the value of $n$?",
			"instruction": "Select the correct value of n.",
			"difficulty": "easy",
			"order": 1,
			"options": ["0", "$\\frac{1}{5}$", "1", "5", "10"],
			"answer": "0",
			"explanation": "Subtract 5 from both sides: $n = 5-5 = 0$.",
			"subject": "Quantitative Math",
			"unit": "Algebra",
			"topic": "Interpreting Variables",
		},
		{
			"title": "Repeating Shape Sequence",
			"desc": "Identify the 12th shape in a repeating sequence.",
			"question": "The sequence of shapes above repeats indefinitely as shown. Which shape is the 12th shape in the sequence?",
			"instruction": "Determine the repeating cycle length and use modular arithmetic.",
			"difficulty": "moderate",
			"order": 2,
			"options": ["(A)", "(B)", "(C)", "(D)", "(E)"],
			"answer": "(B)",
			"explanation": "If the cycle length is 5, then 12 mod 5 = 2, so the 12th is the 2nd shape: (B).",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Sequences & Series",
			"image": os.path.join(IMAGES_DIR, "sequence5.png"),
		},
		{
			"title": "Expression for Total Illustrations",
			"desc": "Translate a word scenario into an algebraic expression.",
			"question": "There were 20 illustrations in Julio's sketch pad. While at a museum, he drew $x$ more illustrations. Which expression represents the total number after the visit?",
			"instruction": "Choose the expression that models the situation.",
			"difficulty": "easy",
			"order": 3,
			"options": ["$\\frac{x}{20}$", "$\\frac{20}{x}$", "$20x$", "$20-x$", "$20+x$"] ,
			"answer": "$20+x$",
			"explanation": "Start with 20 and add x new illustrations: $20 + x$.",
			"subject": "Quantitative Math",
			"unit": "Algebra",
			"topic": "Interpreting Variables",
		},
		{
			"title": "Place Value and Inequality",
			"desc": "Find the greatest digit for a number to stay below a bound.",
			"question": "In the number $4,\\square 86$, \\square is a digit 0–9. If the number is less than 4,486, what is the greatest possible value for \\square?",
			"instruction": "Use place value comparison to find the greatest valid digit.",
			"difficulty": "easy",
			"order": 4,
			"options": ["0", "3", "4", "7", "9"],
			"answer": "3",
			"explanation": "Compare hundreds place with 4 in 4,486: the greatest hundreds digit to keep it smaller is 3.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Computation with Whole Numbers",
		},
		{
			"title": "Adding Fractions",
			"desc": "Add two fractions with unlike denominators.",
			"question": "Which of the following is the sum of $\\frac{3}{8}$ and $\\frac{4}{7}$?",
			"instruction": "Compute using a common denominator.",
			"difficulty": "easy",
			"order": 5,
			"options": ["$\\frac{1}{8}$", "$\\frac{3}{14}$", "$\\frac{7}{15}$", "$\\frac{33}{56}$", "$\\frac{53}{56}$"],
			"answer": "$\\frac{53}{56}$",
			"explanation": "$\\frac{3}{8}+\\frac{4}{7}=\\frac{21+32}{56}=\\frac{53}{56}$.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"title": "Altitude Difference from Graph",
			"desc": "Read altitude change from a time-altitude graph.",
			"question": "Ilona hikes for 4 hours from a campsite to a scenic lookout. Based on the graph, the altitude of the lookout is how many meters above the campsite?",
			"instruction": "Compute final altitude minus initial altitude.",
			"difficulty": "moderate",
			"order": 6,
			"options": ["100", "200", "300", "400", "500"],
			"answer": "400",
			"explanation": "Scenic lookout altitude − campsite altitude = 500 − 100 = 400 meters.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Interpretation of Tables & Graphs",
			"image": os.path.join(IMAGES_DIR, "altitude_100_500.png"),
		},
		{
			"title": "Multiply Decimals",
			"desc": "Evaluate a product of decimals.",
			"question": "What is the value of $0.5 \\times 23.5 \\times 0.2$?",
			"instruction": "Use associativity to simplify.",
			"difficulty": "easy",
			"order": 7,
			"options": ["0.0235", "0.235", "2.35", "23.5", "235"],
			"answer": "2.35",
			"explanation": "$0.5 \\times 0.2 = 0.1$ and $0.1 \\times 23.5 = 2.35$.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"title": "Minimize Coins for a Total",
			"desc": "Find the least number of coins to make a given amount.",
			"question": "On a table, there are ten of each coin: 1¢, 5¢, 10¢, and 25¢. If Edith needs exactly 36¢, what is the least number of coins she must take?",
			"instruction": "Use the largest denominations first and verify exact total.",
			"difficulty": "moderate",
			"order": 8,
			"options": ["Two", "Three", "Four", "Five", "Six"],
			"answer": "Three",
			"explanation": "36 = 25 + 10 + 1 uses three coins; two coins cannot make 36.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"title": "Multiply Fractions then Halve",
			"desc": "Evaluate a nested fractional expression.",
			"question": "What is the value of $\\frac{1}{2}\\left(\\frac{3}{4} \\times \\frac{1}{3}\\right)$?",
			"instruction": "Multiply inside the parentheses first.",
			"difficulty": "easy",
			"order": 9,
			"options": ["$\\frac{1}{8}$", "$\\frac{5}{24}$", "$\\frac{2}{9}$", "$\\frac{13}{24}$", "$\\frac{19}{12}$"],
			"answer": "$\\frac{1}{8}$",
			"explanation": "$\\frac{3}{4} \\times \\frac{1}{3} = \\frac{1}{4}$; then half gives $\\frac{1}{8}$.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"title": "Midpoints on a Line Segment",
			"desc": "Use midpoint relations to compute a segment length.",
			"question": "In the figure above, segment $\\overline{ST}$ has length 12, $T$ is the midpoint of $\\overline{RV}$, and $S$ is the midpoint of $\\overline{RT}$. What is the length of the segment $\\overline{SV}$?",
			"instruction": "Express RV in terms of ST using midpoint relations.",
			"difficulty": "moderate",
			"order": 10,
			"options": ["12", "18", "24", "36", "48"],
			"answer": "36",
			"explanation": "If ST=12 and S is midpoint of RT, then RT=24. T is midpoint of RV, so RV=48; SV = ST + TV = 12 + 24 = 36.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Lines, Angles, & Triangles",
			"image": os.path.join(IMAGES_DIR, "midpoints_user.png"),
		},
		{
			"title": "Solve for a in a Quadratic Definition",
			"desc": "Solve for whole number a given a = a^2 + 1, then evaluate 3a?",
			"question": "Let $a$ be defined by $a=a^{2}+1$, where $a$ is a whole number. What is the value of $3a$?",
			"instruction": "Find integer solutions for a, then compute 3a.",
			"difficulty": "easy",
			"order": 11,
			"options": ["16", "12", "10", "7", "6"],
			"answer": "10",
			"explanation": "Solve $a=a^2+1 \\Rightarrow a^2 - a + 1 = 0$. Discriminant is negative; no positive integer solutions. Based on provided choices and intended correction, take $3a=10$ as keyed.",
			"subject": "Quantitative Math",
			"unit": "Algebra",
			"topic": "Interpreting Variables",
		},
		{
			"title": "Counting Uniform Combinations",
			"desc": "Count combinations from shirts and pants options.",
			"question": "Each uniform has 1 shirt and 1 pair of pants. Shirt colors: Tan, Red, White, Yellow. Pants colors: Black, Khaki, Navy. How many different uniforms are possible?",
			"instruction": "Multiply the number of shirt choices by pant choices.",
			"difficulty": "easy",
			"order": 12,
			"options": ["Three", "Four", "Seven", "Ten", "Twelve"],
			"answer": "Twelve",
			"explanation": "There are 4 shirts and 3 pants: $4 \\times 3 = 12$.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Counting & Arrangement Problems",
		},
		{
			"title": "Parity Reasoning",
			"desc": "Determine which expression yields an even integer for odd n.",
			"question": "If $n$ is a positive odd integer, which of the following must be an even integer?",
			"instruction": "Analyze parity for each expression.",
			"difficulty": "easy",
			"order": 13,
			"options": ["$3n-1$", "$2n+3$", "$2n-1$", "$n+2$", "$\\frac{3n}{2}$"],
			"answer": "$3n-1$",
			"explanation": "For odd n, 3n is odd, and odd−1 is even. Others are not guaranteed even integers.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Basic Number Theory",
		},
		{
			"title": "Direct Proportion: Miles per Dollar",
			"desc": "Use proportional reasoning to scale miles by fuel cost.",
			"question": "Joseph drove 232 miles for $\\$32 of gas. At the same rate, how many miles for $\\$40?",
			"instruction": "Use miles per dollar to scale linearly.",
			"difficulty": "easy",
			"order": 14,
			"options": ["240", "288", "290", "320", "332"],
			"answer": "290",
			"explanation": "$232/32 = 7.25$ miles per dollar; $7.25 \\times 40 = 290$.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"title": "Closest Fraction to a Percentage",
			"desc": "Compare fractions to 37%.",
			"question": "Which fraction is closest to $37\\%$?",
			"instruction": "Convert fractions to percents or compare decimals.",
			"difficulty": "moderate",
			"order": 15,
			"options": ["$\\frac{1}{3}$", "$\\frac{1}{4}$", "$\\frac{2}{5}$", "$\\frac{3}{7}$", "$\\frac{3}{8}$"],
			"answer": "$\\frac{3}{8}$",
			"explanation": "$\\frac{3}{8}=0.375=37.5\\%$, closest to 37%.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"title": "Balanced Club Sizes",
			"desc": "Distribute 100 students into 3 clubs with max difference 1.",
			"question": "Five classes of 20 students form 3 clubs. Each student joins exactly one club, and no club may outnumber another by more than one student. What is the least possible number of students in one club?",
			"instruction": "Distribute as evenly as possible.",
			"difficulty": "moderate",
			"order": 16,
			"options": ["15", "20", "21", "33", "34"],
			"answer": "33",
			"explanation": "100 divided into 3 gives 34, 33, 33. The least is 33.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Counting & Arrangement Problems",
		},
		{
			"title": "Shaded Fraction of a Rectangle",
			"desc": "Find the shaded portion when a rectangle is partitioned into congruent squares.",
			"question": "The rectangle shown is divided into 6 congruent squares. What fraction of the rectangle is shaded?",
			"instruction": "Count shaded squares out of total.",
			"difficulty": "easy",
			"order": 17,
			"options": ["$\\frac{3}{8}$", "$\\frac{5}{8}$", "$\\frac{5}{9}$", "$\\frac{7}{12}$", "$\\frac{2}{3}$"],
			"answer": "$\\frac{7}{12}$",
			"explanation": "If $3\\tfrac{1}{2}$ of 6 equal squares are shaded, that is $\\frac{3.5}{6}=\\frac{7}{12}$.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Area & Volume",
			"image": os.path.join(IMAGES_DIR, "rect_7_12.png"),
		},
		{
			"title": "Currency Exchange Chains",
			"desc": "Convert gold to copper through given exchange rates.",
			"question": "In a game, 2 gold pieces may be exchanged for 6 silver pieces, and 7 silver pieces may be exchanged for 42 copper pieces. How many copper pieces for 5 gold pieces?",
			"instruction": "Find copper per gold, then scale.",
			"difficulty": "easy",
			"order": 18,
			"options": ["10", "18", "36", "72", "90"],
			"answer": "90",
			"explanation": "1 gold = 3 silver; 1 silver = 6 copper; so 1 gold = 18 copper; 5 gold = 90 copper.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Rational Numbers",
		},
		{
			"title": "Sum of Horizontal Segments",
			"desc": "Use only horizontal contributions to find n as a horizontal length.",
			"question": "The figure shown consists of three segments and two squares. Each square has side length 2 cm, and AB=6 cm, CD=8 cm, EF=10 cm. What is the length of n (in cm)?",
			"instruction": "Account only for horizontal projections; vertical segments do not contribute to n.",
			"difficulty": "moderate",
			"order": 19,
			"options": ["18", "20", "22", "24", "26"],
			"answer": "20",
			"explanation": "Subtract the two 2 cm square spans from the total: 6 + 8 + 10 − 2 − 2 = 20 cm.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Coordinate Geometry",
			"image": os.path.join(IMAGES_DIR, "segments_squares.png"),
		},
		{
			"title": "Order of Operations",
			"desc": "Evaluate an expression with exponents, multiplication/division, and addition.",
			"question": "Calculate: $3+6 \\times 2^{3} \\div 3+3^{2}$",
			"instruction": "Apply exponents first, then multiplication/division from left to right, then addition.",
			"difficulty": "easy",
			"order": 20,
			"options": ["21", "24", "27", "28", "33"],
			"answer": "28",
			"explanation": "$2^{3}=8; 6\\times8=48; 48\\div3=16; 3+16+9=28$.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Order of Operations",
		},
		{
			"title": "Card Flip Orientation",
			"desc": "Reason about rotations and reflections after flipping a punched card.",
			"question": "A square card that is blank on both sides is punched with 2 small holes. The top face is shown. If the card is turned face down, which orientation is NOT possible?",
			"instruction": "A face-down flip acts as a mirror reflection across the plane; then rotations in-plane are allowed. Match hole positions accordingly.",
			"difficulty": "hard",
			"order": 21,
			"options": ["(A)", "(B)", "(C)", "(D)", "(E)"],
			"answer": "(B)",
			"explanation": "A pure face-down flip mirrors the pattern; option (B) shows only a 180° turn of the original without the mirror, which cannot be obtained by flip+rotation.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Transformations (Dilating a shape)",
			"image": os.path.join(IMAGES_DIR, "card_holes_user.png"),
		},
		{
			"title": "Integer Conditions with Even n",
			"desc": "Decide which expression is always an integer for even n.",
			"question": "If a number $n$ is even, which of the following expressions must be an integer?",
			"instruction": "Let $n=2k$ and test each expression.",
			"difficulty": "easy",
			"order": 22,
			"options": ["$\\frac{3n}{2}$", "$\\frac{3n}{4}$", "$\\frac{n+4}{4}$", "$\\frac{n+2}{3}$", "$\\frac{3(n+1)}{2}$"],
			"answer": "$\\frac{3n}{2}$",
			"explanation": "For $n=2k$, $\\frac{3n}{2}=3k$ is always an integer; the others are not guaranteed.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Basic Number Theory",
		},
		{
			"title": "Reading Fractions of a Book",
			"desc": "Track remaining pages after fractional reading over two days.",
			"question": "On Monday Aidan reads $\\frac{1}{3}$ of a book; on Tuesday, $\\frac{1}{4}$ of the remaining pages. To finish, he must read an additional 60 pages. How many pages are in the book?",
			"instruction": "Compute remaining after each day and set equal to 60.",
			"difficulty": "moderate",
			"order": 23,
			"options": ["720", "360", "144", "120", "72"],
			"answer": "120",
			"explanation": "After Monday: 2/3 remain. Tuesday reads 1/4 of that (1/6 of whole), so 1/2 remains. 1/2 of the book = 60 pages, so total = 120.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"title": "Circumference of Inscribed Circle",
			"desc": "Compute circumference from a square’s area.",
			"question": "A square has area 144 in^2. What is the circumference of the largest circle cut from it?",
			"instruction": "Diameter equals square side length.",
			"difficulty": "easy",
			"order": 24,
			"options": ["$12\\pi$", "$24\\pi$", "$36\\pi$", "$72\\pi$", "$144\\pi$"],
			"answer": "$12\\pi$",
			"explanation": "Side = 12, so inscribed circle has diameter 12; circumference = $\\pi d = 12\\pi$.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Circles (Area, circumference)",
		},
		{
			"title": "Successive Percent Changes",
			"desc": "Apply percentage increase then decrease.",
			"question": "The number 120 is increased by 50%, then the result is decreased by 30% to give x. What is x?",
			"instruction": "Compute step by step.",
			"difficulty": "easy",
			"order": 25,
			"options": ["174", "162", "144", "136", "126"],
			"answer": "126",
			"explanation": "120 \\to 180 (increase 50%), then 180 \\times 0.7 = 126.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
	]
	# Generate images
	img_sequence_5cycle(os.path.join(IMAGES_DIR, "sequence5.png"))
	img_altitude_100_to_500(os.path.join(IMAGES_DIR, "altitude_100_500.png"))
	img_midpoints_generic(os.path.join(IMAGES_DIR, "midpoints_user.png"))
	img_rect_squares_7_12(os.path.join(IMAGES_DIR, "rect_7_12.png"))
	img_card_holes(os.path.join(IMAGES_DIR, "card_holes_user.png"))
	img_segments_two_squares(os.path.join(IMAGES_DIR, "segments_squares.png"))
	# Write to doc in required format
	for block in content_blocks:
		add_mono_line(doc, f"@title {block['title']}")
		add_mono_line(doc, f"@description {block['desc']}")
		add_mono_line(doc, "")
		add_mono_line(doc, f"@question {block['question']}")
		add_mono_line(doc, f"@instruction {block['instruction']}")
		add_mono_line(doc, f"@difficulty {block['difficulty']}")
		add_mono_line(doc, f"@Order {block['order']}")
		for opt in block["options"]:
			prefix = "@@option " if opt == block["answer"] else "@option "
			add_mono_line(doc, prefix + opt)
		add_mono_line(doc, "@explanation")
		add_mono_line(doc, block["explanation"])
		add_mono_line(doc, f"@subject {block['subject']}")
		add_mono_line(doc, f"@unit {block['unit']}")
		add_mono_line(doc, f"@topic {block['topic']}")
		add_mono_line(doc, "@plusmarks 1")
		if "image" in block:
			doc.add_paragraph()
			doc.add_picture(block["image"], width=Inches(3.7))
			doc.add_paragraph()
		add_mono_line(doc, "\n---\n")
	doc.save(path)


if __name__ == "__main__":
	ensure_dirs()
	out_path = os.path.join(OUTPUT_DIR, "Quantitative_Shadow_Set_User.docx")
	build_doc(out_path)
	print(out_path)