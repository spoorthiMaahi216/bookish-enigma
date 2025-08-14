import os
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw
import math

OUTPUT_DIR = "/workspace/shadow_questions"
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")


def ensure_dirs() -> None:
	os.makedirs(IMAGES_DIR, exist_ok=True)


def add_mono_line(doc: Document, text: str) -> None:
	p = doc.add_paragraph()
	run = p.add_run(text)
	font = run.font
	font.name = "Courier New"
	font.size = Pt(10.5)


# ---------- Image generators ----------

def img_sequence(path: str) -> None:
	w, h = 680, 130
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	for i in range(8):
		x = 15 + i * 82
		y = 20
		t = i % 4
		if t == 0:
			d.ellipse((x, y, x + 58, y + 58), outline="black", width=3)
		elif t == 1:
			d.rectangle((x, y, x + 58, y + 58), outline="black", width=3)
		elif t == 2:
			d.polygon([(x + 29, y), (x + 58, y + 58), (x, y + 58)], outline="black")
		else:
			cx, cy, r = x + 29, y + 29, 27
			pts = []
			for k in range(5):
				ang = -math.pi / 2 + k * 2 * math.pi / 5
				pts.append((cx + r * math.cos(ang), cy + r * math.sin(ang)))
			star = [pts[i % 5] for i in [0, 2, 4, 1, 3]]
			d.line(star + [star[0]], fill="black", width=3)
	img.save(path)


def img_midpoints(path: str) -> None:
	w, h = 660, 120
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	margin = 30
	R = (margin, h // 2)
	S = (margin + 120, h // 2)
	T = (margin + 240, h // 2)
	V = (margin + 480, h // 2)
	d.line((R, V), fill="black", width=3)
	for label, pt in [("R", R), ("S", S), ("T", T), ("V", V)]:
		d.ellipse((pt[0] - 4, pt[1] - 4, pt[0] + 4, pt[1] + 4), fill="black")
		d.text((pt[0] - 6, pt[1] - 24), label, fill="black")
	img.save(path)


def img_rect_squares(path: str) -> None:
	w, h = 390, 260
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	cell_w, cell_h = w // 3, h // 2
	shaded = {(0, 0), (1, 0), (2, 1)}
	for r in range(2):
		for c in range(3):
			x0, y0 = c * cell_w, r * cell_h
			x1, y1 = x0 + cell_w - 2, y0 + cell_h - 2
			if (c, r) in shaded:
				d.rectangle((x0 + 2, y0 + 2, x1, y1), fill=(185, 185, 185))
			d.rectangle((x0 + 2, y0 + 2, x1, y1), outline="black", width=3)
	img.save(path)


def img_altitude(path: str) -> None:
	w, h = 520, 280
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	# axes
	d.line((50, 20, 50, h - 40), fill="black", width=2)
	d.line((50, h - 40, w - 20, h - 40), fill="black", width=2)
	# simple upward polyline from (0,200) to (4,550)
	points = [
		(50, h - 40 - 0),  # origin visually
	]
	# map time 0..4 to x, altitude 200..550 to y
	for t, alt in [(0, 200), (1, 300), (2, 370), (3, 460), (4, 550)]:
		x = 50 + int(t * (w - 100) / 4)
		y = h - 40 - int((alt - 200) * (h - 80) / (550 - 200))
		points.append((x, y))
	d.line(points[1:], fill="blue", width=3)
	img.save(path)


def img_circle_in_square(path: str) -> None:
	w, h = 280, 280
	img = Image.new("RGB", (w, h), "white")
	d = ImageDraw.Draw(img)
	d.rectangle((20, 20, w - 20, h - 20), outline="black", width=3)
	d.ellipse((20, 20, w - 20, h - 20), outline="black", width=3)
	img.save(path)


def build_questions():
	# Each entry strictly uses topics from provided curriculum
	return [
		{
			"q": "If n − 7 = 5, what is the value of n?",
			"instr": "Solve for n.",
			"difficulty": "easy",
			"order": 1,
			"opts": ["10", "12", "−2", "2", "7"],
			"ans": "12",
			"exp": "Add 7 to both sides: n = 5 + 7 = 12.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Computation with Whole Numbers",
		},
		{
			"q": "The shapes repeat every 4 (circle, square, triangle, star). Which shape is 12th?",
			"instr": "Use modular arithmetic.",
			"difficulty": "easy",
			"order": 2,
			"opts": ["Circle", "Square", "Triangle", "Star", "Hexagon"],
			"ans": "Star",
			"exp": "12 mod 4 = 0, so it is the 4th shape in the cycle: Star.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Sequences & Series",
			"image": os.path.join(IMAGES_DIR, "sequence.png"),
		},
		{
			"q": "A box has 15 marbles. You add x more. Which expression is the total?",
			"instr": "Model with a variable.",
			"difficulty": "easy",
			"order": 3,
			"opts": ["15/x", "x/15", "15x", "15 − x", "15 + x"],
			"ans": "15 + x",
			"exp": "Start with 15 and add x, giving 15 + x.",
			"subject": "Quantitative Math",
			"unit": "Algebra",
			"topic": "Interpreting Variables",
		},
		{
			"q": "The number 4,□32 is less than 4,532. What is the greatest possible value of □?",
			"instr": "Compare by place value.",
			"difficulty": "easy",
			"order": 4,
			"opts": ["2", "3", "4", "5", "9"],
			"ans": "4",
			"exp": "Hundreds digit must be < 5; the greatest such digit is 4.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Basic Number Theory",
		},
		{
			"q": "What is the sum of 5/12 and 7/18?",
			"instr": "Find a common denominator.",
			"difficulty": "moderate",
			"order": 5,
			"opts": ["11/36", "17/30", "29/36", "41/36", "5/30"],
			"ans": "29/36",
			"exp": "LCM(12,18)=36; 5/12=15/36 and 7/18=14/36; sum = 29/36.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"q": "A hiker starts at 200 m and ends at 550 m after 4 hours. What is the altitude gain?",
			"instr": "Compute the difference.",
			"difficulty": "easy",
			"order": 6,
			"opts": ["100", "200", "300", "350", "400"],
			"ans": "350",
			"exp": "Final − initial = 550 − 200 = 350 m.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Interpretation of Tables & Graphs",
			"image": os.path.join(IMAGES_DIR, "altitude.png"),
		},
		{
			"q": "What is 0.4 × 12.5 × 0.2?",
			"instr": "Multiply stepwise.",
			"difficulty": "easy",
			"order": 7,
			"opts": ["0.1", "1.0", "0.8", "0.5", "0.04"],
			"ans": "1.0",
			"exp": "0.4×12.5=5; 5×0.2=1.0.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"q": "Using 1c, 5c, 10c, and 25c coins (ten of each available), least number of coins to make 37 cents?",
			"instr": "Minimize the count of coins.",
			"difficulty": "moderate",
			"order": 8,
			"opts": ["2", "3", "4", "5", "6"],
			"ans": "4",
			"exp": "25 + 10 + 1 + 1 uses 4 coins.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"q": "What is (1/2) × (2/3 × 3/4)?",
			"instr": "Simplify inside first.",
			"difficulty": "easy",
			"order": 9,
			"opts": ["1/8", "1/4", "1/3", "1/2", "2/3"],
			"ans": "1/4",
			"exp": "(2/3×3/4)=1/2; then 1/2×1/2=1/4.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"q": "On line RV, T is the midpoint of RV and S is the midpoint of RT. If ST = 10, what is SV?",
			"instr": "Use midpoint ratios.",
			"difficulty": "moderate",
			"order": 10,
			"opts": ["10", "15", "20", "30", "40"],
			"ans": "30",
			"exp": "ST = (1/4)RV ⇒ RV=40. Then SV = RV − RS = 40 − 10 = 30.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Lines, Angles, & Triangles",
			"image": os.path.join(IMAGES_DIR, "midpoints.png"),
		},
		{
			"q": "Let b be a nonzero whole number such that b = b^2 − 2b. What is b?",
			"instr": "Solve the quadratic and use the nonzero condition.",
			"difficulty": "moderate",
			"order": 11,
			"opts": ["0", "1", "2", "3", "4"],
			"ans": "3",
			"exp": "b=b^2−2b ⇒ 0=b^2−3b ⇒ b(b−3)=0. Nonzero b gives b=3.",
			"subject": "Quantitative Math",
			"unit": "Algebra",
			"topic": "Quadratic Equations & Functions (Finding roots/solutions, graphing)",
		},
		{
			"q": "A uniform consists of 1 shirt and 1 pair of pants. If there are 4 shirt colors and 3 pants colors, how many different uniforms are possible?",
			"instr": "Apply the multiplication principle.",
			"difficulty": "easy",
			"order": 12,
			"opts": ["7", "10", "12", "24", "36"],
			"ans": "12",
			"exp": "4 × 3 = 12.",
			"subject": "Quantitative Math",
			"unit": "Data Analysis & Probability",
			"topic": "Counting & Arrangement Problems",
		},
		{
			"q": "If n is a positive odd integer, which of the following must be an even integer?",
			"instr": "Test each option with parity rules.",
			"difficulty": "easy",
			"order": 13,
			"opts": ["3n − 1", "2n + 3", "2n − 1", "n + 2", "(3n)/2"],
			"ans": "3n − 1",
			"exp": "Odd×3 = odd, and odd−1 is even. Others are odd or not guaranteed integers.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Basic Number Theory",
		},
		{
			"q": "A car travels 180 miles using $27 of gasoline. At the same rate, how many miles for $40?",
			"instr": "Use direct proportion.",
			"difficulty": "easy",
			"order": 14,
			"opts": ["240", "260", "267", "280", "300"],
			"ans": "267",
			"exp": "Miles per dollar = 180/27 = 6.666…; ×40 ≈ 266.7 ≈ 267.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"q": "Which fraction is closest to 41%?",
			"instr": "Compare decimal values.",
			"difficulty": "moderate",
			"order": 15,
			"opts": ["1/3", "2/5", "3/7", "3/8", "5/12"],
			"ans": "5/12",
			"exp": "5/12 ≈ 41.67% is closest to 41%.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
		{
			"q": "There are 100 students forming 3 clubs with sizes that differ by at most 1. What is the least possible club size?",
			"instr": "Distribute as evenly as possible.",
			"difficulty": "moderate",
			"order": 16,
			"opts": ["30", "31", "32", "33", "34"],
			"ans": "33",
			"exp": "100/3 ≈ 33.33 ⇒ sizes 33, 33, 34; least is 33.",
			"subject": "Quantitative Math",
			"unit": "Problem Solving",
			"topic": "Problem Solving",
		},
		{
			"q": "A rectangle is divided into 6 congruent squares; 3 are shaded as shown. What fraction is shaded?",
			"instr": "Count shaded squares over total squares.",
			"difficulty": "easy",
			"order": 17,
			"opts": ["1/2", "3/5", "3/6", "2/3", "5/6"],
			"ans": "1/2",
			"exp": "3 of the 6 equal squares are shaded.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Area & Volume",
			"image": os.path.join(IMAGES_DIR, "rect_squares.png"),
		},
		{
			"q": "If 3 gold = 12 silver and 4 silver = 28 copper, how many copper for 5 gold?",
			"instr": "Convert units step by step.",
			"difficulty": "moderate",
			"order": 18,
			"opts": ["35", "40", "56", "70", "140"],
			"ans": "140",
			"exp": "1 gold = 4 silver; 1 silver = 7 copper ⇒ 1 gold = 28 copper ⇒ 5 gold = 140 copper.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Rational Numbers",
		},
		{
			"q": "A straight bar is formed by segments of 6 cm, 8 cm, and 10 cm placed end-to-end. Two square caps of side 2 cm are attached at the two joints, extending length by one side each. What is the total length n (cm)?",
			"instr": "Add the segment lengths and the two added side lengths.",
			"difficulty": "moderate",
			"order": 19,
			"opts": ["24", "26", "28", "30", "32"],
			"ans": "28",
			"exp": "6 + 8 + 10 + 2 + 2 = 28 cm.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Perimeter",
		},
		{
			"q": "Compute: 5 + (8 × 2^3 ÷ 4) + 2^2",
			"instr": "Follow order of operations.",
			"difficulty": "easy",
			"order": 20,
			"opts": ["19", "21", "23", "25", "27"],
			"ans": "25",
			"exp": "2^3=8 ⇒ 8×8/4=64/4=16; then 5+16+4=25.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Order of Operations",
		},
		{
			"q": "A figure is reflected across a vertical line. Which transformation describes this?",
			"instr": "Identify the transformation type.",
			"difficulty": "easy",
			"order": 21,
			"opts": ["Rotation", "Translation", "Reflection", "Dilation", "Shear"],
			"ans": "Reflection",
			"exp": "Flipping across a line is a reflection.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Transformations (Dilating a shape)",
		},
		{
			"q": "If n is even, which expression must be an integer?",
			"instr": "Let n = 2k and test.",
			"difficulty": "easy",
			"order": 22,
			"opts": ["(n + 2)/2", "(3n)/4", "(n + 1)/2", "(n + 6)/4", "(3n + 3)/2"],
			"ans": "(n + 2)/2",
			"exp": "n=2k ⇒ (n+2)/2 = (2k+2)/2 = k+1, always an integer.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Rational Numbers",
		},
		{
			"q": "On Monday, Aidan reads 1/4 of a book. On Tuesday, he reads 1/5 of the remaining pages. He has 90 pages left. How many pages are in the book?",
			"instr": "Track the remaining fraction.",
			"difficulty": "moderate",
			"order": 23,
			"opts": ["150", "180", "225", "300", "360"],
			"ans": "150",
			"exp": "Remaining after Monday: 3/4. After Tuesday: (4/5)(3/4) = 3/5. If 3/5 = 90, total = 150.",
			"subject": "Quantitative Math",
			"unit": "Reasoning",
			"topic": "Word Problems",
		},
		{
			"q": "A square has area 196 square inches. What is the circumference of the largest inscribed circle?",
			"instr": "Relate side length to circle diameter.",
			"difficulty": "moderate",
			"order": 24,
			"opts": ["14π", "28π", "42π", "56π", "196π"],
			"ans": "14π",
			"exp": "Side = √196 = 14 = diameter ⇒ circumference = πd = 14π.",
			"subject": "Quantitative Math",
			"unit": "Geometry and Measurement",
			"topic": "Circles (Area, circumference)",
			"image": os.path.join(IMAGES_DIR, "circle_in_square.png"),
		},
		{
			"q": "A number 200 is increased by 20% and then decreased by 25% to give x. What is x?",
			"instr": "Use successive multipliers.",
			"difficulty": "easy",
			"order": 25,
			"opts": ["140", "150", "160", "170", "180"],
			"ans": "180",
			"exp": "1.20 × 0.75 = 0.9 ⇒ 200 × 0.9 = 180.",
			"subject": "Quantitative Math",
			"unit": "Numbers and Operations",
			"topic": "Fractions, Decimals, & Percents",
		},
	]


def build_doc(path: str) -> None:
	doc = Document()
	add_mono_line(doc, "@title Quantitative Reasoning Shadow Set A")
	add_mono_line(doc, "@description 25 MCQ shadow questions inspired by provided base set with images where applicable")
	questions = build_questions()
	# Generate images needed
	img_sequence(os.path.join(IMAGES_DIR, "sequence.png"))
	img_midpoints(os.path.join(IMAGES_DIR, "midpoints.png"))
	img_rect_squares(os.path.join(IMAGES_DIR, "rect_squares.png"))
	img_altitude(os.path.join(IMAGES_DIR, "altitude.png"))
	img_circle_in_square(os.path.join(IMAGES_DIR, "circle_in_square.png"))
	# Add questions
	for item in questions:
		add_mono_line(doc, "@question " + item["q"]) 
		add_mono_line(doc, "@instruction " + item["instr"]) 
		add_mono_line(doc, "@difficulty " + item["difficulty"]) 
		add_mono_line(doc, f"@Order {item['order']}")
		for opt in item["opts"]:
			prefix = "@@option " if opt == item["ans"] else "@option "
			add_mono_line(doc, prefix + opt)
		add_mono_line(doc, "@explanation ")
		add_mono_line(doc, item["exp"]) 
		add_mono_line(doc, "@subject " + item["subject"]) 
		add_mono_line(doc, "@unit " + item["unit"]) 
		add_mono_line(doc, "@topic " + item["topic"]) 
		add_mono_line(doc, "@plusmarks 1")
		if "image" in item:
			doc.add_paragraph()
			doc.add_picture(item["image"], width=Inches(3.5))
			doc.add_paragraph()
		doc.add_paragraph()
	# Save
	doc.save(path)


if __name__ == "__main__":
	ensure_dirs()
	out_path = os.path.join(OUTPUT_DIR, "Quantitative_Shadow_Set_A.docx")
	build_doc(out_path)
	print(out_path)